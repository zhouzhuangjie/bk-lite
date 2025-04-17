import re
import tempfile

import docx
from langchain_core.documents import Document
from tqdm import tqdm

from src.ocr.base_ocr import BaseOCR


class DocLoader:
    def __init__(self, file_path: str, ocr: BaseOCR, mode: str = 'full'):
        """
        mode: full, paragraph
        """

        self.file_path = file_path
        self.mode = mode
        self.ocr = ocr

    def table_to_md(self, table):
        # Converts a docx table to markdown format
        md_table = []
        for row in table.rows:
            md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
            md_table.append(md_row)
        return '\n'.join(md_table)

    def remove_unicode_chars(self, text):
        return re.sub(r'\\u[fF]{1}[0-9a-fA-F]{3}', '', text)

    def load(self):
        docs = []
        try:
            document = docx.Document(self.file_path)
        except Exception as e:
            raise IOError(f"Error loading document: {e}")

        paragraphs = document.paragraphs

        if self.mode == 'full':
            full_text = ""
            for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
                full_text += paragraph.text
            docs.append(Document(full_text))

        elif self.mode == 'paragraph':
            current_doc = None
            for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
                if any(heading in paragraph.style.name for heading in ('Heading', '标题')):
                    if current_doc is not None:
                        docs.append(Document(current_doc.strip()))
                    current_doc = paragraph.text.strip() + "\n"  # Start a new
                else:
                    if current_doc is not None:
                        current_doc += paragraph.text.strip() + "\n"
                    else:
                        current_doc = paragraph.text.strip() + "\n"

            if current_doc:
                docs.append(Document(current_doc.strip()))

        else:
            raise ValueError("Invalid mode.")

        tables = document.tables
        for table in tqdm(tables, desc=f"解析[{self.file_path}]的表格"):
            docs.append(Document(self.table_to_md(table),
                        metadata={"format": "table"}))

        # 提取图片并使用OCR识别
        if self.ocr is not None:
            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    try:
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as temp_img:
                            temp_img.write(image_data)
                            temp_img.flush()
                            ocr_result = self.ocr.predict(temp_img.name)
                            docs.append(
                                Document(ocr_result, metadata={"format": "image"}))
                    except Exception as e:
                        raise IOError(f"Error processing image: {e}")

        return docs
