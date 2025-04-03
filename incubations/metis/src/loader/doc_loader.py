import re

import docx
from langchain_core.documents import Document
from tqdm import tqdm


class DocLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def table_to_md(self, table):
        # Converts a docx table to markdown format
        md_table = []
        for row in table.rows:
            md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
            md_table.append(md_row)
        return '\n'.join(md_table)

    def remove_unicode_chars(self, text):
        return re.sub(r'\\u[fF]{1}[0-9a-fA-F]{3}', '', text)

    def load(self):
        docs = []
        try:
            document = docx.Document(self.file_path)
        except Exception as e:
            raise IOError(f"Error loading document: {e}")

        paragraphs = document.paragraphs

        full_text = ""
        for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
            full_text += paragraph.text
        docs.append(Document(full_text))

        # Process tables
        tables = document.tables
        for table in tqdm(tables, desc=f"解析[{self.file_path}]的表格"):
            docs.append(Document(self.table_to_md(table), metadata={"format": "table"}))

        return docs
